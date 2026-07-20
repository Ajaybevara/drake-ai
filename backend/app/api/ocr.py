from __future__ import annotations

import asyncio
import base64
import uuid
from io import BytesIO
from pathlib import Path
from typing import Literal

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from PIL import Image

from app.core.config import settings
from app.core.security import get_current_user
from app.models import User

router = APIRouter()

OCR_ROOT = Path("uploads") / "ocr"
MAX_UPLOAD_BYTES = 150 * 1024 * 1024
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
_rapid_ocr_engine = None


class ExportRequest(BaseModel):
    text: str
    format: Literal["docx", "pdf"]


def _save_upload(user_id: int, filename: str, data: bytes) -> str:
    folder = OCR_ROOT / str(user_id)
    folder.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(ch if ch.isalnum() or ch in "._- " else "-" for ch in filename).strip() or "ocr-file"
    target = folder / f"{uuid.uuid4().hex[:8]}-{safe_name[:90]}"
    target.write_bytes(data)
    return str(target)


def _image_to_data_url(image: Image.Image) -> str:
    if image.mode != "RGB":
        image = image.convert("RGB")
    buffered = BytesIO()
    image.save(buffered, format="JPEG", quality=92)
    encoded = base64.b64encode(buffered.getvalue()).decode("utf-8")
    return f"data:image/jpeg;base64,{encoded}"


async def _vllm_ocr_available() -> tuple[bool, str]:
    if not settings.VLLM_OCR_BASE_URL:
        return False, "VLLM_OCR_BASE_URL is not configured"
    try:
        from openai import AsyncOpenAI

        client = AsyncOpenAI(api_key=settings.VLLM_API_KEY, base_url=settings.VLLM_OCR_BASE_URL)
        await asyncio.wait_for(client.models.list(), timeout=8)
        return True, ""
    except Exception as exc:
        return False, str(exc)


def _local_ocr_available() -> bool:
    try:
        import rapidocr_onnxruntime  # noqa: F401

        return True
    except Exception:
        return False


def _get_local_ocr_engine():
    global _rapid_ocr_engine
    if _rapid_ocr_engine is None:
        try:
            from rapidocr_onnxruntime import RapidOCR

            _rapid_ocr_engine = RapidOCR()
        except Exception as exc:
            raise HTTPException(status_code=503, detail=f"Local OCR engine is not available: {exc}") from exc
    return _rapid_ocr_engine


def _extract_text_with_local_ocr(image: Image.Image) -> str:
    try:
        import numpy as np

        if image.mode != "RGB":
            image = image.convert("RGB")
        engine = _get_local_ocr_engine()
        result, _ = engine(np.array(image))
        lines: list[str] = []
        for item in result or []:
            if len(item) < 2:
                continue
            text = str(item[1]).strip()
            score = float(item[2]) if len(item) > 2 and item[2] is not None else 1.0
            if text and score >= 0.35:
                lines.append(text)
        return "\n".join(lines).strip()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=503, detail=f"Local OCR extraction failed: {exc}") from exc


async def _extract_text_from_image(image: Image.Image) -> tuple[str, str]:
    reachable, _ = await _vllm_ocr_available()
    if reachable:
        return await _extract_text_with_vllm(image), "vllm-vision-ocr"
    text = _extract_text_with_local_ocr(image)
    if not text:
        raise HTTPException(status_code=422, detail="OCR completed, but no readable text was found in this image")
    return text, "local-rapidocr"


async def _extract_text_with_vllm(image: Image.Image) -> str:
    reachable, error = await _vllm_ocr_available()
    if not reachable:
        raise HTTPException(
            status_code=503,
            detail=(
                "VLLM OCR server is not reachable. Start the Qwen2.5-VL OCR server at "
                f"{settings.VLLM_OCR_BASE_URL}. Original error: {error}"
            ),
        )
    try:
        from openai import AsyncOpenAI

        client = AsyncOpenAI(api_key=settings.VLLM_API_KEY, base_url=settings.VLLM_OCR_BASE_URL)
        prompt = (
            "You are an advanced OCR assistant. Extract all visible text from this image with high accuracy. "
            "Preserve document structure, headings, paragraphs, tables, numbers, symbols, and units. "
            "If handwriting or unclear text is present, transcribe it as accurately as possible and mark uncertain text with [?]. "
            "Return only clean Markdown. Do not add explanations that are not present in the image."
        )
        response = await asyncio.wait_for(
            client.chat.completions.create(
                model=settings.VLLM_OCR_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": _image_to_data_url(image)}},
                        ],
                    }
                ],
                temperature=0,
                max_tokens=4096,
            ),
            timeout=settings.VLLM_TIMEOUT_SECONDS,
        )
        text = (response.choices[0].message.content or "").strip()
        return text.replace("```markdown", "").replace("```", "").strip()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=503, detail=f"VLLM OCR failed: {exc}") from exc


def _render_pdf_pages(data: bytes, max_pages: int = 80) -> list[Image.Image]:
    try:
        import fitz

        pages: list[Image.Image] = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page_index in range(min(len(doc), max_pages)):
                page = doc.load_page(page_index)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.open(BytesIO(pix.tobytes("png")))
                pages.append(image)
        return pages
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not render PDF pages: {exc}") from exc


def _extract_embedded_pdf_text(data: bytes) -> str:
    try:
        import fitz

        with fitz.open(stream=data, filetype="pdf") as doc:
            parts = []
            for idx, page in enumerate(doc):
                text = page.get_text("text").strip()
                if text:
                    parts.append(f"## Page {idx + 1}\n\n{text}")
            return "\n\n".join(parts)
    except Exception:
        return ""


def _word_bytes(text: str) -> bytes:
    try:
        from docx import Document

        doc = Document()
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped:
                doc.add_paragraph(stripped)
            else:
                doc.add_paragraph("")
        output = BytesIO()
        doc.save(output)
        return output.getvalue()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {exc}") from exc


def _pdf_bytes(text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        output = BytesIO()
        doc = SimpleDocTemplate(output, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        flowables = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                flowables.append(Spacer(1, 8))
                continue
            style = styles["Normal"]
            content = stripped
            if stripped.startswith("# "):
                style = styles["Heading1"]
                content = stripped[2:]
            elif stripped.startswith("## "):
                style = styles["Heading2"]
                content = stripped[3:]
            elif stripped.startswith("### "):
                style = styles["Heading3"]
                content = stripped[4:]
            elif stripped.startswith("- "):
                content = f"• {stripped[2:]}"
            flowables.append(Paragraph(content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style))
            flowables.append(Spacer(1, 8))
        doc.build(flowables)
        return output.getvalue()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"PDF export failed: {exc}") from exc


@router.get("/status")
async def ocr_status(current_user: User = Depends(get_current_user)):
    configured = bool(settings.VLLM_OCR_BASE_URL)
    reachable, error = await _vllm_ocr_available()
    local_ready = _local_ocr_available()
    return {
        "mode": "vllm-vision-ocr" if reachable else ("local-rapidocr" if local_ready else "local-pdf-text-fallback"),
        "vllm_configured": configured,
        "vllm_reachable": reachable,
        "vllm_ocr_base_url": settings.VLLM_OCR_BASE_URL,
        "vllm_ocr_model": settings.VLLM_OCR_MODEL,
        "local_ocr_available": local_ready,
        "error": error,
    }


@router.post("/extract-image")
async def extract_image(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="No file uploaded")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is larger than 150 MB")
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in IMAGE_SUFFIXES:
        raise HTTPException(status_code=400, detail="Upload PNG, JPG, TIFF, BMP, or WEBP image files")
    try:
        image = Image.open(BytesIO(data))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read image: {exc}") from exc
    storage_path = _save_upload(current_user.id, file.filename or "ocr-image.jpg", data)
    text, mode = await _extract_text_from_image(image)
    return {"text": text, "pages": 1, "filename": file.filename, "storage_path": storage_path, "mode": mode}


@router.post("/extract-pdf")
async def extract_pdf(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="No file uploaded")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is larger than 150 MB")
    if Path(file.filename or "").suffix.lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Upload a PDF file")
    storage_path = _save_upload(current_user.id, file.filename or "ocr-document.pdf", data)

    embedded = _extract_embedded_pdf_text(data)
    if embedded:
        return {"text": embedded, "pages": embedded.count("## Page "), "filename": file.filename, "storage_path": storage_path, "mode": "local-pdf-text-fallback"}

    reachable, error = await _vllm_ocr_available()
    if reachable:
        pages = _render_pdf_pages(data)
        extracted = []
        for idx, image in enumerate(pages):
            page_text = await _extract_text_with_vllm(image)
            extracted.append(f"## Page {idx + 1}\n\n{page_text}")
        return {"text": "\n\n".join(extracted), "pages": len(pages), "filename": file.filename, "storage_path": storage_path, "mode": "vllm-vision-ocr"}

    if _local_ocr_available():
        pages = _render_pdf_pages(data)
        extracted = []
        for idx, image in enumerate(pages):
            page_text = _extract_text_with_local_ocr(image)
            extracted.append(f"## Page {idx + 1}\n\n{page_text or '[No readable text found on this page]'}")
        return {"text": "\n\n".join(extracted), "pages": len(pages), "filename": file.filename, "storage_path": storage_path, "mode": "local-rapidocr"}

    raise HTTPException(status_code=503, detail=f"OCR engine is unavailable. Start vLLM at {settings.VLLM_OCR_BASE_URL}. Original error: {error}")


@router.post("/export")
def export_text(req: ExportRequest, current_user: User = Depends(get_current_user)):
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="No extracted text to export")
    if req.format == "docx":
        content = _word_bytes(req.text)
        filename = "drake_ocr_extracted_text.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        content = _pdf_bytes(req.text)
        filename = "drake_ocr_extracted_text.pdf"
        mime = "application/pdf"
    return JSONResponse({
        "content": base64.b64encode(content).decode("utf-8"),
        "filename": filename,
        "mime": mime,
    })
